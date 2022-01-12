'''

    22/01/12

Sub fit()
Worksheets("Sheet1").Columns("A:AN").AutoFit
End Sub

https://stackoverflow.com/questions/24023518/using-python-to-autofit-all-columns-of-an-excel-sheet

from win32com.client import Dispatch

excel = Dispatch('Excel.Application')
wb = excel.Workbooks.Open("D:\\output.xlsx")

#Activate second sheet
excel.Worksheets(2).Activate()

#Autofit column in active sheet
excel.ActiveSheet.Columns.AutoFit()

#Save changes in a new file
wb.SaveAs("D:\\output_fit.xlsx")

#Or simply save changes in a current file
#wb.Save()

wb.Close()

'''
import re
import os
import calendar
import time
from docx import Document
from docx.shared import Inches
import schedule
from docx.shared import Inches, Cm
import xlsxwriter as xlw
from datetime import date, timedelta
import bwxrefcom
    
_week_days = ['월','화','수','목','금','토', '일']
_sunday = 6
_days_per_page = 150
_ncolumn_sub = 3
_max_read_day = 300
_bible_reading_verse_data = "300.dat"

def add_table_item_title(table, ncol):
    for i in range(ncol):
        table.rows[0].cells[i*4+0].text = '일차'
        table.rows[0].cells[i*4+1].text = '날짜'
        table.rows[0].cells[i*4+2].text = '요일'
        table.rows[0].cells[i*4+3].text = '오늘의 말씀'
    
def access_denied(e_str):
    key = ["access", "denied", "used", "another", "permission"]
    return any(x in e_str.lower() for x in key)
    
def auto_fit_excel_column(file):
    from win32com.client import Dispatch
    
    try:
        excel = Dispatch('Excel.Application')
        wb = excel.Workbooks.Open(file)
    
        #Activate first sheet
        excel.Worksheets(1).Activate()
        
        #Autofit column in active sheet
        excel.ActiveSheet.Columns.AutoFit()
        
        #Save changes in a new file
        #wb.SaveAs("D:\\output_fit.xlsx")
        
        #Or simply save changes in a current file
        wb.Save()
        wb.Close()
    except Exception as e:
        e_str = str(e)
        if access_denied(e_str):
            e_str += "\n%s is already opened!"%file
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
    finally:
        excel.Application.Quit()

def create_bible_reading_schedule_excel(
        file_out, 
        year,
        start_mon, 
        end_mon, 
        ncol,
        auto_fit, 
        delay_time,
        include_sunday=False):
        
    #print(xlsx_file, year, start_mon, end_mon, ncol, auto_fit)
    ncolumn_sub = 4
    try:
        workbook = xlw.Workbook(file_out)
    except Exception as e:
        e_str = "... Error: Can't create Excel Document.\n%s"%str(e)
        if access_denied(e_str):
            e_str += "\n%s is already opened!"%file_out
            bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return False
        
    worksheet = workbook.add_worksheet()
    checkbox_format = workbook.add_format({'bold': True, 'font_size': 9, 'font_color': 'black'})
    cell_format = workbook.add_format({'font_size': 8})
        
    month_days = list(map(lambda mon: calendar.monthrange(year,mon), range(start_mon,end_mon+1)))
    
    nrow, mod = divmod(_max_read_day, ncol)
    
    i_total_day = 1
    i_day_per_month = 0
    i_day_per_page = 0
    i_week = 1
    i_row = 0
    i_col = 0
    i_day = 0
    days_per_page = nrow*ncol
    
    for i_mon, m_data in enumerate(month_days):
        j_mon = i_mon + start_mon
        #print(calendar.month_name[j_mon])
        first_day_of_month = m_data[0] # the first day of the month 
        total_day_of_month = m_data[1]
        i_day_of_month = 0
        i_weekday = first_day_of_month
        
        while i_day_of_month < total_day_of_month:
            if not include_sunday and i_weekday == _sunday: 
                i_day_of_month += 1
                i_weekday = 0
                continue
            elif include_sunday and i_weekday > _sunday:
                i_weekday = 0

            if i_day >= _max_read_day:
                break
                
            info = schedule.table[i_day]
            #date = "%s/%2d/%2d (%s)"% (str(year)[-2:],j_mon,i_day_of_month+1, _week_days[i_weekday])
            date = "%2d/%2d (%s)"% (j_mon,i_day_of_month+1, _week_days[i_weekday])
            worksheet.write(i_row, i_col*ncolumn_sub+0, info[0], cell_format)
            worksheet.write(i_row, i_col*ncolumn_sub+1, date, cell_format)
            worksheet.write(i_row, i_col*ncolumn_sub+2, info[1], cell_format)
            worksheet.write(i_row, i_col*ncolumn_sub+3, '□', checkbox_format)
            
            i_day += 1
            i_row += 1
            i_weekday += 1
            i_day_of_month += 1
            i_day_per_page += 1
            if i_row >= nrow:
                i_row = 0
                i_col += 1
            
            #if i_day_per_page >= days_per_page:
            #    i_row = 0
            #    i_col = 0
            #    i_day_per_page = 0
                #document.add_page_break()
                #table = document.add_table(rows, cols)
    try:
        workbook.close()
    except Exception as e:
        e_str = "... Error: Can't create Excel Document.\n%s"%str(e)
        if access_denied(e_str):
            e_str += "\n%s is already opened!"%file_out
            bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return False
        
    if auto_fit:
        time.sleep(delay_time)
        auto_fit_excel_column(os.path.join(os.getcwd(), file_out))
    #print('success')
    bwxrefcom.message_box(bwxrefcom.message_normal, "Success")
    
def create_bible_reading_schedule_word(
        file_out, 
        year,
        start_mon, 
        end_mon, 
        nrow, 
        ncol,
        include_sunday=False):

    #print(file_out, year, start_mon, end_mon, nrow, ncol)
    
    try:
        document = Document('default.docx')
    except Exception as e:
        e_str = "... Error(xref_to_docx): Can't create Word Document.\n%s"%str(e)
        bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return False
        
    sections = document.sections
    half_inch = 2.54*0.5
    for section in sections:
        section.top_margin    = Cm(half_inch)
        section.bottom_margin = Cm(half_inch)
        section.left_margin   = Cm(half_inch)
        section.right_margin  = Cm(half_inch)

    # Tuple (start day, total days)
    # start day: mon(0) - sun(6)
    # total days
    month_days = list(map(lambda mon: calendar.monthrange(year,mon), range(start_mon,end_mon+1)))
    #
    # 150 + 150
    # 150/3 = 50
    rows = nrow
    cols = ncol*_ncolumn_sub
    #msg.appendPlainText('... Table(row,col): %d x %d'%(rows,cols))
    table = document.add_table(rows, cols)
    days_per_page = nrow * ncol
    #
    #  Week  Date(M-S)  Book(chap/verse)
    #
    #add_table_item_title(table)
    
    i_total_day = 1
    i_day_per_month = 0
    i_day_per_page = 0
    i_week = 1
    i_row = 0
    i_col = 0
    i_day = 0
    days_per_page = nrow*ncol
    
    for i_mon, m_data in enumerate(month_days):
        j_mon = i_mon + start_mon
        print(calendar.month_name[j_mon])
        first_day_of_month = m_data[0] # the first day of the month 
        total_day_of_month = m_data[1]
        i_day_of_month = 0
        i_weekday = first_day_of_month
        
        while i_day_of_month < total_day_of_month:
            if not include_sunday and i_weekday == _sunday: 
                i_day_of_month += 1
                i_weekday = 0
                continue
            elif include_sunday and i_weekday > _sunday:
                i_weekday = 0

            if i_day >= _max_read_day:
                break
                
            info = schedule.table[i_day]
            table.rows[i_row].cells[i_col*_ncolumn_sub+0].text = info[0]
            table.rows[i_row].cells[i_col*_ncolumn_sub+1].text = "%s/%2d/%2d (%s)"%\
            (str(year)[-2:],j_mon,i_day_of_month+1, _week_days[i_weekday])
            table.rows[i_row].cells[i_col*_ncolumn_sub+2].text = info[1]
            i_day += 1
            i_row += 1
            i_weekday += 1
            i_day_of_month += 1
            i_day_per_page += 1
            if i_row >= nrow:
                i_row = 0
                i_col += 1
            
            if i_day_per_page >= days_per_page:
                i_row = 0
                i_col = 0
                i_day_per_page = 0
                document.add_page_break()
                table = document.add_table(rows, cols)
                
            #if include_sunday and i_weekday > _sunday: 
            #    #i_day_of_month += 1
            #    i_weekday = 0
            #    #continue
          
    try:
        document.save(file_out)
    except Exception as e:
        e_str = "... Error: Can't create Excel Document.\n%s"%str(e)
        if access_denied(e_str):
            e_str += "\n%s is already opened!"%file_out
            bwxrefcom.message_box(bwxrefcom.message_error, e_str)
        return False
        
    bwxrefcom.message_box(bwxrefcom.message_normal, "Success")
    
def find_last_date(mon, day, year, days, include_start_date=True):
    date1 = date(year, mon, day)
    date2 = date1 + timedelta(days=days-1 if include_start_date else days)
    return date2
    
#create_bible_reading_schedule_word('test.docx', 2022, 2, 2, 25, 2, True)
#xlsx_file = 'test.xlsx'
#create_bible_reading_schedule_excel(xlsx_file, 2, 13, 8, 2022, True)
#auto_fit_excel_column(os.path.join(os.getcwd(), xlsx_file))
#print(find_last_date(2,1,2022, 300))